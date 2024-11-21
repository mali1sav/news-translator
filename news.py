import streamlit as st
import os
import asyncio
import httpx
import json
from dotenv import load_dotenv
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO
from typing import List, Dict, Any
from docx import Document

# Load environment variables
load_dotenv()

# Constants
OPENROUTER_API_KEY = os.getenv("OPENROUTER_API_KEY")
if not OPENROUTER_API_KEY:
    st.error("OPENROUTER_API_KEY not found in environment variables.")
    st.stop()

API_BASE_URL = "https://openrouter.ai/api/v1"
HEADERS = {"Authorization": f"Bearer {OPENROUTER_API_KEY}"}

# Initialize Async Client
client = httpx.AsyncClient(base_url=API_BASE_URL, headers=HEADERS)

# Revised TAGS Structure
TAGS = {
    "Primary Categories": [
        "Bitcoin News",
        "Ethereum News",
        "Altcoin News",
        "DeFi News",
        "NFT News",
        "Industry News"
    ],
    "Topics": [
        "Market Trends",
        "Exchange News",
        "Technical Analysis",
        "Cyber Security",
        "Technology",
        "Regulation",
        "Adoption",
        "Investment",
        "Partnership",
        "SEC",
        "Legal"
    ],
    "Cryptocurrencies": {
        "Layer 1": [
            "Solana",
            "Cardano",
            "Avalanche",
            "Polkadot",
            "Cosmos",
            "NEAR Protocol",
            "Fantom",
            "TRON",
            "Other Layer 1"

        ],
        "Layer 2": [
            "Polygon",
            "Arbitrum",
            "Optimism",
            "Base",
            "Other Layer 2"
        ],
        "Infrastructure": [
            "Chainlink",
            "The Graph",
            "Quant",
            "Ren Protocol",
            "Other Infrastructure"
        ],
        "DeFi": [
            "Protocols",
            "Uniswap",
            "Aave",
            "Maker",
            "Curve",
            "Synthetix",
            "Compound",
            "Other DeFi platforms"
        ],
        "AI": [
            "Render",
            "Fetch.ai",
            "Ocean Protocol",
            "TAO",
            "Other AI coins"
        ],

        "Gaming & Metaverse": [
            "The Sandbox",
            "Decentraland",
            "Axie Infinity",
            "Immutable X",
            "Other Gaming & Metaverse"
        ],
        "Meme": [
            "Dogecoin",
            "Shiba Inu",
            "Pepe",
            "Floki",
            "Bonk",
            "Friend.tech",
            "Memecoin",
            "Other Meme"
        ],
        "Stable Coins": [
            "Tether",
            "USDC",
            "DAI",
            "Other Stable Coins"
        ]
    },
    "Market Infrastructure": {
        "Exchanges": [
            "Coinbase",
            "Binance",
            "Kraken",
            "KuCoin",
            "OKx",
            "Gemini",
            "Bitfinex",
            "Bitstamp",
            "Other Exchanges"
        ],
        "Data & Analytics": [
            "CoinGecko",
            "CoinMarketCap",
            "Glassnode",
            "Messari",
            "Chainalysis",
            "Etherscan",
            "Other Data & Analytics"
        ],
        "Wallets & Custody": [
            "Ledger",
            "Trezor",
            "MetaMask",
            "Trust Wallet",
            "Exodus",
            "Tangem",
            "Other Wallets & Custody"
        ]
    }
}

# Update the ALL_TAGS generation
ALL_TAGS = set(TAGS["Primary Categories"])
ALL_TAGS.update(TAGS["Topics"])
for category in TAGS["Cryptocurrencies"].values():
    ALL_TAGS.update(category)
for category in TAGS["Market Infrastructure"].values():
    ALL_TAGS.update(category)
ALL_TAGS = sorted(ALL_TAGS)  # Sorted list for consistent ordering

async def generate_meta_description(content: str) -> str:
    """
    Generate a Thai meta description using LLM.
    """
    system_prompt = """
You are a professional crypto content writer. Create a meta description in Thai for this crypto news article.
Requirements:
- Keep length STRICTLY no more than 120 characters
- Capture the key points: price movements, market impact, and key events
- Use natural Thai language
- Focus on the main news angle and impact
- If contain names such as people, company, or coin names, DO NOT translate names but ensure to translate the rest
- Return the meta description in the following JSON format:
{
    "meta_description": "your Thai meta description here"
}
"""
    try:
        response = await client.post(
            "/chat/completions",
            json={
                "model": "openai/gpt-4o-2024-11-20",
                "messages": [
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": f"Create a Thai meta description for this article:\n{content}"}
                ],
                "response_format": {
                    "type": "json_object"
                }
            },
            timeout=60
        )
        response.raise_for_status()
        data = response.json()
        
        # Parse the response content as JSON
        content = data['choices'][0]['message']['content']
        parsed_response = json.loads(content)
        thai_meta = parsed_response['meta_description'].strip()
        
        # Ensure it's not longer than 120 characters
        if len(thai_meta) > 120:
            thai_meta = thai_meta[:117] + "..."
        return thai_meta
    except Exception as e:
        st.error(f"Meta description generation error: {str(e)}")
        return "ไม่มีคำอธิบาย"

async def translate_text(text: str, is_title: bool = False, is_h1: bool = False) -> str:
    """
    Translate text from English to Thai using OpenRouter's translation service with structured output.
    """
    if is_title:
        system_prompt = """You are a Thai crypto news translator. Translate this title to Thai.
Rules:
- Create a compelling news headline
- Maximum 60 Thai characters. 
- Do not truncate the content; instead, use concise wording to fit within the limit.
- Keep cryptocurrency names in English
- Keep other entity and technical terms in English (e.g. people names, organisation names, platforms, technical concepts, product names etc.) throughout the article, even when the surrounding text is in Thai.
- Use English names for cryptocurrencies instead of their Thai transliterations throughout the article. For example, use 'Bitcoin' instead of 'บิทคอยน์', 'Ethereum' instead of 'อีเธอเรียม', 'Solana' instead of 'โซลาน่า', etc.
- Return the translation in the following JSON format:
{
    "translated_text": "your Thai translation here"
}"""
    elif is_h1:
        system_prompt = """You are a Thai crypto news translator. Translate this H1 to Thai.
Rules:
- Create a news-like headline between 8-15 words
- Keep cryptocurrency names in English
- Align with Title.
- Keep other entity and technical terms in English (e.g. people names, organisation names, platforms, technical concepts, product names etc.) throughout the article, even when the surrounding text is in Thai.
- Use English names for cryptocurrencies instead of their Thai transliterations throughout the article. For example, use 'Bitcoin' instead of 'บิทคอยน์', 'Ethereum' instead of 'อีเธอเรียม', 'Solana' instead of 'โซลาน่า', etc.
- Return the translation in the following JSON format:
{
    "translated_text": "your Thai translation here"
}"""
    else:
        system_prompt = """You are a Thai crypto news translator. Translate this text to Thai.
Rules:
- Use natural Thai language
- Write for Thai audiences with basic crypto knowledge, using semi-professional Thai language
- Use Thai crypto terms where appropriate
- Maintain the original meaning while making it natural in Thai
- Keep other entity and technical terms in English (e.g. people names, organisation names, platforms, technical concepts, product names etc.) throughout the article, even when the surrounding text is in Thai.
- Use English names for cryptocurrencies instead of their Thai transliterations throughout the article. For example, use 'Bitcoin' instead of 'บิทคอยน์', 'Ethereum' instead of 'อีเธอเรียม', 'Solana' instead of 'โซลาน่า', etc.
- Return the translation in the following JSON format:
{
    "translated_text": "your Thai translation here"
}"""

    try:
        response = await client.post(
            "/chat/completions",
            json={
                "model": "openai/gpt-4o-2024-11-20",
                "messages": [
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": text}
                ],
                "response_format": {
                    "type": "json_object"
                }
            },
            timeout=60
        )
        response.raise_for_status()
        data = response.json()
        
        # Parse the response content as JSON
        content = data['choices'][0]['message']['content']
        parsed_response = json.loads(content)
        translated_text = parsed_response['translated_text'].strip()
        
        # Additional check for title length
        if is_title and len(translated_text) > 70:
            # Retry with stronger emphasis on length
            return await translate_text(text, is_title=True)
            
        return translated_text
        
    except Exception as e:
        st.error(f"Translation error: {str(e)}")
        return text

async def categorize_with_llm(content: str) -> List[str]:
    """
    Categorize the news article using Gemini and return a list of relevant tags.
    """
    system_prompt = """You are a cryptocurrency news categorization expert. Analyze this article and suggest relevant tags.

Content Analysis Requirements:
1. Primary Category (Must include one):
   - If about Bitcoin: "Bitcoin News"
   - If about Ethereum: "Ethereum News"
   - If about other cryptocurrencies: "Altcoin News"

2. Specific Cryptocurrencies:
   - Identify all mentioned cryptocurrencies
   - Include their ecosystem tags (e.g., "Solana", "Cardano", etc.)

3. Market Infrastructure:
   - Identify mentioned exchanges, wallets, or data providers
   - Include relevant infrastructure tags

4. Topics:
   - Market analysis/trends
   - Technical developments
   - Regulatory news
   - Business developments

5. Special Categories:
   - Check for DeFi, NFT, or Gaming content
   - Identify if content relates to Layer 1/Layer 2 solutions

Rules:
- Maximum 5 most relevant tags
- Only use tags from the provided list
- Prioritize specificity over generality
- Ensure tags accurately reflect the main focus of the article

Return only comma-separated tags from the available list."""

    try:
        response = await client.post(
            "/chat/completions",
            json={
                "model": "google/gemini-flash-1.5-8b",
                "messages": [
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": f"Available tags:\n{', '.join(ALL_TAGS)}\n\nArticle content for analysis:\n{content}"}
                ],
                "temperature": 0.3,  # Lower temperature for more focused responses
                "max_tokens": 100    # Limit response length
            },
            timeout=60
        )
        response.raise_for_status()
        data = response.json()
        tags_str = data['choices'][0]['message']['content'].strip()
        
        # Process and validate tags
        suggested_tags = [tag.strip() for tag in tags_str.split(',') if tag.strip() in ALL_TAGS]
        
        # Ensure we have at least one primary category
        has_primary = any(tag in TAGS["Primary Categories"] for tag in suggested_tags)
        if not has_primary:
            # Add most relevant primary category based on content
            if 'bitcoin' in content.lower():
                suggested_tags.insert(0, "Bitcoin News")
            elif 'ethereum' in content.lower():
                suggested_tags.insert(0, "Ethereum News")
            else:
                suggested_tags.insert(0, "Altcoin News")
        
        return suggested_tags[:5]  # Return top 5 tags
        
    except Exception as e:
        st.error(f"Categorization error: {e}")
        return []

async def translate_section(section: Dict[str, str]) -> Dict[str, str]:
    """
    Translate a single section's heading and content.
    """
    translated_heading = await translate_text(section.get('heading', ''))
    translated_content = await translate_text(section.get('content', ''))
    return {
        'heading': translated_heading,
        'content': translated_content
    }

def parse_article(content: str) -> Dict[str, Any]:
    """
    Parse the article content into structured elements with content split into chunks.
    Each chunk contains 2-3 paragraphs separated by blank lines.
    """
    try:
        parsed_content = {
            'title': '',
            'meta_description': '',
            'h1': '',
            'sections': []
        }

        # Split content into lines and clean up
        lines = [line.strip() for line in content.split('\n') if line.strip()]

        # Skip unwanted content
        skip_patterns = [
            'Crypto Reporter',
            'Share',
            'Last updated:',
            'Why Trust Cryptonews',
            'GMT'
        ]

        content_lines = [line for line in lines if not any(pattern in line for pattern in skip_patterns)]

        if not content_lines:
            raise ValueError("No valid content found after filtering.")

        # Extract title (first line)
        parsed_content['title'] = content_lines[0]
        parsed_content['h1'] = content_lines[0]

        # Assume the first paragraph after the title is the meta description
        if len(content_lines) > 1:
            parsed_content['meta_description'] = content_lines[1]
        else:
            parsed_content['meta_description'] = "Not provided"

        # Process main content into sections based on headings
        current_section = {'heading': 'Main Content', 'content': []}
        for line in content_lines[1:]:
            if line.isupper() or line.endswith(':'):  # Simple heuristic for headings
                if current_section['content']:
                    current_section['content'] = '\n\n'.join(current_section['content'])
                    parsed_content['sections'].append(current_section)
                current_section = {'heading': line, 'content': []}
            else:
                current_section['content'].append(line)
        # Add the last section
        if current_section['content']:
            current_section['content'] = '\n\n'.join(current_section['content'])
            parsed_content['sections'].append(current_section)

        # Further split main content into chunks of 2-3 paragraphs
        for section in parsed_content['sections']:
            paragraphs = section['content'].split('\n\n')
            chunks = [paragraphs[i:i + 3] for i in range(0, len(paragraphs), 3)]
            chunked_sections = []
            for chunk in chunks:
                chunked_sections.append({
                    'heading': section['heading'],
                    'content': '\n\n'.join(chunk)
                })
            section['chunks'] = chunked_sections
            del section['content']  # Remove original content

        # Validate parsed content
        if not parsed_content['title'] or not parsed_content['sections']:
            raise ValueError("Missing required content elements")

        return parsed_content

    except Exception as e:
        st.error(f"Error parsing article content: {str(e)}")
        return {'title': '', 'meta_description': 'Not provided', 'h1': '', 'sections': []}

async def translate_content(structured_content: Dict[str, Any]) -> Dict[str, Any]:
    """
    Translate each element of the structured content and categorize with tags.
    """
    translated = {}
    
    # Translate title and h1
    translated['title'] = await translate_text(structured_content.get('title', ''), is_title=True)
    translated['h1'] = await translate_text(structured_content.get('h1', ''), is_h1=True)
    
    # Generate meta description in Thai (no English meta description needed)
    combined_content = f"{structured_content.get('title', '')} {structured_content.get('h1', '')} " \
                    + ' '.join([
                        chunk['content'] 
                        for section in structured_content.get('sections', []) 
                        for chunk in section.get('chunks', [])
                    ])
    translated['meta_description'] = await generate_meta_description(combined_content)
    
    # Translate sections concurrently with progress tracking
    sections = structured_content.get('sections', [])
    translated_sections = []
    
    for section_idx, section in enumerate(sections):
        chunks = section.get('chunks', [])
        translated_chunks = []
        
        # Create progress bar for chunks
        chunk_progress = st.progress(0)
        st.write(f"Translating section {section_idx + 1}/{len(sections)}...")
        
        for chunk_idx, chunk in enumerate(chunks):
            try:
                translated_chunk = await translate_section(chunk)
                translated_chunks.append(translated_chunk)
                # Update progress
                chunk_progress.progress((chunk_idx + 1) / len(chunks))
            except Exception as e:
                st.error(f"Error translating chunk {chunk_idx + 1} in section {section_idx + 1}: {e}")
                # Add empty translation to maintain structure
                translated_chunks.append({
                    'heading': chunk.get('heading', ''),
                    'content': 'Translation error occurred'
                })
        
        translated_sections.append({
            'heading': section.get('heading', ''),
            'chunks': translated_chunks
        })
        
        # Clear progress bar after section is done
        chunk_progress.empty()
    
    translated['sections'] = translated_sections
    
    # Get tags from LLM
    tags = await categorize_with_llm(combined_content)
    
    # Add mandatory tags based on content analysis
    mandatory_tags = []
    content_lower = combined_content.lower()
    
    # Check for major cryptocurrencies
    if 'bitcoin' in content_lower or ' btc ' in content_lower:
        mandatory_tags.extend(['Bitcoin News', 'BTC'])
    if 'ethereum' in content_lower or ' eth ' in content_lower:
        mandatory_tags.extend(['Ethereum News', 'ETH'])
    if 'solana' in content_lower or ' sol ' in content_lower:
        mandatory_tags.extend(['Altcoin News', 'SOL'])
    
    # Check for specific entities
    for category, entities in TAGS["Market Infrastructure"].items():
        for entity in entities:
            if entity.lower() in content_lower:
                if entity not in mandatory_tags:
                    mandatory_tags.append(entity)
    
    # Remove duplicates while preserving order
    seen = set()
    all_tags = []
    for tag in (mandatory_tags + tags):
        if tag not in seen and tag in ALL_TAGS:
            seen.add(tag)
            all_tags.append(tag)
    
    translated['tags'] = all_tags[:5]  # Keep top 5 tags
    return translated

def determine_primary_category(crypto: str) -> str:
    """
    Determine the primary category based on the cryptocurrency.
    """
    if crypto.lower() == 'bitcoin':
        return 'Bitcoin News'
    elif crypto.lower() == 'ethereum':
        return 'Ethereum News'
    else:
        return 'Altcoin News'

def create_docx(original: Dict[str, Any], translated: Dict[str, Any]) -> BytesIO:
    """
    Create a properly formatted Word document with structured output.
    """
    doc = Document()
    
    # Set styles for different heading levels and ensure black color
    styles = doc.styles
    for style_name in ['Heading 1', 'Heading 2', 'Heading 3']:
        style = styles[style_name]
        style.font.color.rgb = None  # This will make it black
    
    # Add URL
    url_para = doc.add_paragraph()
    url_run = url_para.add_run('Original URL: ')
    url_run.bold = True
    url_para.add_run(original.get('url', ''))
    
    # Add separator
    doc.add_paragraph('_' * 40)
    
    # Add Title
    title_para = doc.add_paragraph()
    title_label = title_para.add_run('Title: ')
    title_label.bold = True
    title_para.add_run(translated.get('title', ''))
    
    # Add Meta Description
    meta_para = doc.add_paragraph()
    meta_title = meta_para.add_run('Meta Description: ')
    meta_title.bold = True
    meta_para.add_run(translated.get('meta_description', ''))
    
    # Add separator
    doc.add_paragraph('_' * 40)
    
    # Add H1
    h1_para = doc.add_paragraph()
    h1_label = h1_para.add_run('H1: ')
    h1_label.bold = True
    h1_para.add_run(translated.get('h1', ''))
    
    # Add Main Content without headers
    for section in translated.get('sections', []):
        for chunk in section.get('chunks', []):
            if chunk.get('content'):
                para = doc.add_paragraph(chunk.get('content', ''))
                para.paragraph_format.space_after = Pt(12)
    
    # Save to BytesIO
    doc_io = BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    return doc_io

def display_tag_section(current_tags: List[str]):
    """
    Display all tags in 6 columns with clear headings.
    """
    st.markdown("### Tag Management")
    st.write("Current Tags:", ", ".join(current_tags))
    
    # Create 6 columns
    cols = st.columns(6)
    
    # Column 1: Primary Categories
    with cols[0]:
        st.markdown("##### Primary Categories")
        for tag in TAGS["Primary Categories"]:
            key = f"primary_{tag.replace(' ', '_')}"
            if st.checkbox(tag, value=tag in current_tags, key=key):
                if tag not in current_tags:
                    current_tags.append(tag)
                elif tag in current_tags:
                    current_tags.remove(tag)
    
    # Column 2: Topics
    with cols[1]:
        st.markdown("##### Topics")
        for tag in TAGS["Topics"]:
            key = f"topic_{tag.replace(' ', '_')}"
            if st.checkbox(tag, value=tag in current_tags, key=key):
                if tag not in current_tags:
                    current_tags.append(tag)
                elif tag in current_tags:
                    current_tags.remove(tag)
    
    # Column 3-4: Cryptocurrencies (split into two columns due to length)
    crypto_categories = list(TAGS["Cryptocurrencies"].items())
    mid = len(crypto_categories) // 2
    
    with cols[2]:
        st.markdown("##### Cryptocurrencies (1/2)")
        for category, coins in crypto_categories[:mid]:
            st.markdown(f"###### {category}")
            for tag in coins:
                key = f"crypto_{category}_{tag.replace(' ', '_')}"
                if st.checkbox(tag, value=tag in current_tags, key=key):
                    if tag not in current_tags:
                        current_tags.append(tag)
                    elif tag in current_tags:
                        current_tags.remove(tag)
    
    with cols[3]:
        st.markdown("##### Cryptocurrencies (2/2)")
        for category, coins in crypto_categories[mid:]:
            st.markdown(f"###### {category}")
            for tag in coins:
                key = f"crypto_{category}_{tag.replace(' ', '_')}"
                if st.checkbox(tag, value=tag in current_tags, key=key):
                    if tag not in current_tags:
                        current_tags.append(tag)
                    elif tag in current_tags:
                        current_tags.remove(tag)
    
    # Column 5: Market Infrastructure - Exchanges & Data
    with cols[4]:
        st.markdown("##### Market Infrastructure (1/2)")
        for category in ["Exchanges", "Data & Analytics"]:
            st.markdown(f"###### {category}")
            for tag in TAGS["Market Infrastructure"][category]:
                key = f"infra_{category}_{tag.replace(' ', '_').replace('.', '_')}"
                if st.checkbox(tag, value=tag in current_tags, key=key):
                    if tag not in current_tags:
                        current_tags.append(tag)
                    elif tag in current_tags:
                        current_tags.remove(tag)
    
    # Column 6: Market Infrastructure - Wallets & Others
    with cols[5]:
        st.markdown("##### Market Infrastructure (2/2)")
        for category in ["Wallets & Custody"]:
            st.markdown(f"###### {category}")
            for tag in TAGS["Market Infrastructure"][category]:
                key = f"infra_{category}_{tag.replace(' ', '_').replace('.', '_')}"
                if st.checkbox(tag, value=tag in current_tags, key=key):
                    if tag not in current_tags:
                        current_tags.append(tag)
                    elif tag in current_tags:
                        current_tags.remove(tag)
    
    # Update session state
    st.session_state['current_tags'] = current_tags

async def translate_section(section: Dict[str, str]) -> Dict[str, str]:
    """
    Translate a single section's heading and content.
    """
    translated_heading = await translate_text(section.get('heading', ''))
    translated_content = await translate_text(section.get('content', ''))
    return {
        'heading': translated_heading,
        'content': translated_content
    }

def display_comparison(original: Dict[str, Any], translated: Dict[str, Any]):
    """
    Display side-by-side comparison with editable text areas.
    """
    st.markdown("## Content Comparison")

    # Original URL and Tags
    st.text(f"Original URL: {original.get('url', '')}")
    st.text(f"Tags: {', '.join(translated.get('tags', [])) if translated.get('tags') else 'N/A'}")
    st.markdown("---")

    # Initialize session state for edited content if not exists
    if 'edited_content' not in st.session_state:
        st.session_state.edited_content = {
            'title': translated.get('title', ''),
            'h1': translated.get('h1', ''),
            'meta_description': translated.get('meta_description', ''),
            'sections': translated.get('sections', [])
        }

    # Title, H1 & Meta Description
    with st.expander("Title, H1 & Meta Description", expanded=True):
        col1, col2 = st.columns(2)
        with col1:
            st.subheader("English")
            st.text_area("Title", value=original.get('title', ''), height=100, key="en_title", disabled=True)
            st.text_area("H1", value=original.get('h1', ''), height=100, key="en_h1", disabled=True)
            # Empty meta description for English side
            st.text_area("Meta Description", value="", height=100, key="en_meta", disabled=True)
        with col2:
            st.subheader("Thai")
            edited_title = st.text_area("Title", value=st.session_state.edited_content['title'], height=100, key="th_title")
            edited_h1 = st.text_area("H1", value=st.session_state.edited_content['h1'], height=100, key="th_h1")
            edited_meta = st.text_area("Meta Description ภาษาไทย (ไม่เกิน 120 ตัวอักษร)", 
                                     value=st.session_state.edited_content['meta_description'], 
                                     height=100, 
                                     key="th_meta",
                                     help="Meta description should not exceed 120 characters")
            
            # Show character count for meta description
            current_meta_length = len(edited_meta)
            st.write(f"Meta description length: {current_meta_length}/120 characters")
            if current_meta_length > 120:
                st.warning("Meta description exceeds 120 characters limit!")
            
            # Update session state
            st.session_state.edited_content['title'] = edited_title
            st.session_state.edited_content['h1'] = edited_h1
            st.session_state.edited_content['meta_description'] = edited_meta
    # Main Content Sections
    with st.expander("Main Content", expanded=True):
        for i, (orig_section, trans_section) in enumerate(zip(
            original.get('sections', []), 
            translated.get('sections', [])
        )):
            st.markdown(f"### Section {i+1}: {orig_section['heading']}")
            
            # Initialize section in session state if needed
            if i >= len(st.session_state.edited_content['sections']):
                st.session_state.edited_content['sections'].append({
                    'heading': trans_section['heading'],
                    'chunks': trans_section['chunks']
                })
                
            for j, (orig_chunk, trans_chunk) in enumerate(zip(
                orig_section.get('chunks', []), 
                trans_section.get('chunks', [])
            )):
                st.markdown(f"#### Chunk {j+1}")
                col1, col2 = st.columns(2)
                with col1:
                    st.text_area(
                        "English",
                        value=orig_chunk['content'],
                        height=200,
                        key=f"en_section_{i}_chunk_{j}",
                        disabled=True
                    )
                with col2:
                    edited_content = st.text_area(
                        "Thai",
                        value=st.session_state.edited_content['sections'][i]['chunks'][j]['content'],
                        height=200,
                        key=f"th_section_{i}_chunk_{j}"
                    )
                    # Update session state with edited content
                    st.session_state.edited_content['sections'][i]['chunks'][j]['content'] = edited_content
                st.markdown("---")  # Separator between chunks

    # Tags Section
    with st.expander("Tags", expanded=True):
        current_tags = st.session_state.get('current_tags', translated.get('tags', []))
        display_tag_section(current_tags)
    

def process_and_translate(original_url: str, original_article: str):
    """
    Orchestrate the parsing, translation, and categorization of the article.
    """
    # Parse the article
    structured = parse_article(original_article)

    if structured.get('title') and structured.get('meta_description') and structured.get('sections'):
        # Add URL to structured content
        structured['url'] = original_url.strip()

        # Translate the content
        translated = asyncio.run(translate_content(structured))

        # Store in session state
        st.session_state['original'] = structured
        st.session_state['translated'] = translated

        st.success("Content processed and translated successfully.")
    else:
        st.error("Failed to parse the article. Please ensure the format is correct.")

def main():
    try:
        st.set_page_config(layout="wide")
        st.title("Crypto News Translator (EN → TH)")
        
        # Input Section with default values
        st.markdown("### Paste the Original English Article URL and Content Below")

        default_url = "https://cryptonews.com/news/u-s-bitcoin-etfs-surpass-100-billion-in-assets-amid-btc-rally/"
        default_content = """
BlackRock’s IBIT leads the pack with $45.4 billion in assets, followed by Grayscale’s GBTC with $20.6 billion.

U.S. spot Bitcoin exchange-traded funds (ETFs) have surpassed $100 billion in total net assets as Bitcoin continues its record-breaking surge.

Data from SoSoValue shows that as of Wednesday, the 12 spot Bitcoin ETFs collectively held $100.55 billion, accounting for approximately 5.4% of Bitcoin’s total market capitalization.

BlackRock’s IBIT leads the pack with $45.4 billion in assets, followed by Grayscale’s GBTC with $20.6 billion.

Bitcoin Hits New ATH
Bitcoin itself reached an all-time high, trading around $97,094, a 3.8% increase over the past 24 hours.

The funds also experienced significant inflows, with $733.5 million recorded on Wednesday and $837.36 million the day before.

BlackRock’s IBIT saw the highest inflows, receiving $626.5 million, while Fidelity’s FBTC attracted $133.9 million.

Smaller contributions included $9.3 million for Bitwise’s BITB and $3.8 million for Ark and 21Shares’ ARKB. However, Grayscale’s GBTC reported no new inflows.

Trading activity for bitcoin ETFs reached $5.09 billion on Wednesday, down slightly from Tuesday’s $5.71 billion.

Meanwhile, spot Ethereum ETFs in the U.S. continue to see outflows, with $30.3 million withdrawn on Wednesday, marking the fifth consecutive day of negative flows.

Trading volume for ether ETFs dropped to $338.3 million, compared to $345.1 million the previous day.

“The main driver behind BTC’s rapid rise is still institutional involvement. We have seen large net inflows into BTC ETFs this week. By Wednesday this week, BTC ETFs had achieved a net inflow of $1.8 billion,” Gracy Chen, CEO at Bitget, said in a statement.

Chen noted that MicroStrategy purchased 51,000 BTC last week at a cost of $88,617 each, and this week, they announced plans to raise $2.6 billion to continue purchasing BTC.

“Well-known mining companies are planning to issue $850 million in convertible bonds to buy BTC. The massive spot buying power of traditional funds has caused BTC’s price to rise quickly.”

Meanwhile, the open interest in BTC contracts has surged to $63 billion, with a daily increase of $6 billion.

BTC’s implied volatility (IV) has risen to 60, indicating a higher probability of large market fluctuations in the future.

Chen said that short-term capital has a tendency to lock in profits, which could lead to large price swings around the $100K mark.

Bitwise Asset Management Takes Step Towards Solana ETF
Bitwise Asset Management has filed to establish a trust entity for its proposed Bitwise Solana ETF in Delaware, marking an early step toward launching the fund as the crypto ETF market continues to see increased demand.

If approved, Bitwise will join other asset managers like VanEck and 21Shares, which have also pursued Solana-focused ETFs.

The filing follows Bitwise’s recent S-1 registration for an XRP ETF, the first fund proposal offering exposure to Ripple’s native cryptocurrency.

The firm has experienced remarkable growth in 2024, with assets under management (AUM) reaching $5 billion as of October 15—a 400% year-to-date surge.

Bitwise’s spot Bitcoin ETF, BITB, has also gained significant traction, drawing $2.3 billion in net inflows since its launch, second only to offerings by BlackRock and Fidelity.

        """

        original_url = st.text_input(
            label="Original English Article URL",
            value=default_url,
            help="Paste the URL of the original English article here."
        )

        original_article = st.text_area(
            label="Original English Article Content",
            value=default_content,
            height=600,
            key="original_article",
            help="Paste the full English article content here."
        )

        # Process Button
        if st.button("Process and Translate"):
            if not original_url.strip():
                st.warning("Please enter the Original English Article URL.")
            elif not original_article.strip():
                st.warning("Please paste the Original English Article content.")
            else:
                with st.spinner("Processing and translating..."):
                    process_and_translate(original_url, original_article)

        # Display Comparison if available
        if 'original' in st.session_state and 'translated' in st.session_state:
            original = st.session_state['original']
            translated = st.session_state['translated']
            
            # Initialize current_tags in session state if not present
            if 'current_tags' not in st.session_state:
                st.session_state['current_tags'] = translated.get('tags', [])

            # Display content comparison
            display_comparison(original, translated)

            # Create and offer download of the Word document
            if translated:
                doc_io = create_docx(original, st.session_state.edited_content)
                st.download_button(
                    label="Download Edited Thai Content in MS Word",
                    data=doc_io,
                    file_name="translated_article.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

    except Exception as e:
        st.error(f"An error occurred: {str(e)}")

if __name__ == "__main__":
    main()
